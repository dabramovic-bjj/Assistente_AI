import os
import time
import requests
import json
import io
import pandas as pd
from pypdf import PdfReader
from docx import Document
from tavily import TavilyClient
from fpdf import FPDF

# Configurazione
TOKEN = os.environ.get("TELEGRAM_TOKEN", "IL_TUO_TOKEN_TELEGRAM")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "LA_TUA_CHIAVE_OPENAI")
TAVILY_API_KEY = os.environ.get("TAVILY_API_KEY", "TUA_CHIAVE_TAVILY")

tavily = TavilyClient(api_key=TAVILY_API_KEY)
URL = f"https://api.telegram.org/bot{TOKEN}"

chat_histories = {}

def leggi_file_da_telegram(file_id):
    file_info = requests.get(f"{URL}/getFile?file_id={file_id}", timeout=10).json()
    if not file_info.get("ok"): return "Errore nel recupero del file."
    
    file_path = file_info["result"]["file_path"]
    file_url = f"https://api.telegram.org/file/bot{TOKEN}/{file_path}"
    file_bytes = requests.get(file_url, timeout=20).content
    ext = file_path.split(".")[-1].lower()
    
    testo_estratto = ""
    try:
        if ext == "pdf":
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages[:10]:
                testo_estratto += (page.extract_text() or "") + "\n"
        elif ext in ["docx", "doc"]:
            doc = Document(io.BytesIO(file_bytes))
            testo_estratto = "\n".join([p.text for p in doc.paragraphs])
        elif ext in ["xlsx", "xls"]:
            xls = pd.ExcelFile(io.BytesIO(file_bytes))
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name, nrows=50)
                testo_estratto += f"\n--- Foglio: {sheet_name} ---\n" + df.to_string() + "\n"
        elif ext == "csv":
            df = pd.read_csv(io.BytesIO(file_bytes), nrows=50)
            testo_estratto = df.to_string()
        elif ext in ["txt", "py", "json", "md"]:
            testo_estratto = file_bytes.decode("utf-8", errors="ignore")
        else:
            return f"Formato .{ext} non supportato."
    except Exception as e:
        return f"Errore lettura file: {e}"
    return testo_estratto[:15000]

def ask_openai(chat_id, prompt):
    if chat_id not in chat_histories:
        chat_histories[chat_id] = [{"role": "system", "content": "Sei un assistente virtuale utile e intelligente. Oggi è il 21 agosto 2026."}]
    chat_histories[chat_id].append({"role": "user", "content": prompt})
    
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
    data = {"model": "gpt-4o-mini", "messages": chat_histories[chat_id]}
    
    try:
        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=data, timeout=30)
        reply = response.json()["choices"][0]["message"]["content"]
        chat_histories[chat_id].append({"role": "assistant", "content": reply})
        return reply
    except Exception as e:
        return f"Errore: {e}"

def send_message(chat_id, text):
    try:
        requests.post(f"{URL}/sendMessage", json={"chat_id": chat_id, "text": text, "parse_mode": "Markdown"}, timeout=10)
    except: pass

def crea_e_invia_file_modificato(chat_id, testo_modificato, nome_originale):
    ext = nome_originale.split('.')[-1].lower()
    base_name = nome_originale.rsplit('.', 1)[0]
    path = ""
    
    if ext in ['docx', 'doc']:
        doc = Document()
        doc.add_paragraph(testo_modificato)
        path = f"modificato_{base_name}.docx"
        doc.save(path)
    elif ext in ['xlsx', 'xls']:
        try:
            df = pd.read_csv(io.StringIO(testo_modificato))
            path = f"modificato_{base_name}.xlsx"
            df.to_excel(path, index=False)
        except:
            send_message(chat_id, "Errore nella creazione Excel. Assicurati che l'IA abbia fornito dati in formato tabella/CSV.")
            return
    elif ext == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        testo_pulito = testo_modificato.encode('latin-1', 'replace').decode('latin-1')
        for line in testo_pulito.split('\n'):
            pdf.cell(200, 10, txt=line, ln=True)
        path = f"modificato_{base_name}.pdf"
        pdf.output(path)
    
    with open(path, 'rb') as f:
        requests.post(f"{URL}/sendDocument", data={"chat_id": chat_id}, files={"document": f})

def main():
    print("Bot avviato...")
    offset = None
    processed_updates = set()
    while True:
        try:
            updates = requests.get(f"{URL}/getUpdates", params={"timeout": 30, "offset": offset}).json()
            if "result" in updates:
                for update in updates["result"]:
                    offset = update["update_id"] + 1
                    if update["update_id"] in processed_updates: continue
                    processed_updates.add(update["update_id"])
                    
                    if "message" in update:
                        msg = update["message"]
                        chat_id = msg["chat"]["id"]
                        if "document" in msg:
                            file_id = msg["document"]["file_id"]
                            file_name = msg["document"].get("file_name", "documento")
                            istruzioni = msg.get("caption", "")
                            contenuto = leggi_file_da_telegram(file_id)
                            
                            if istruzioni:
                                send_message(chat_id, "✍️ Sto modificando il file...")
                                prompt = f"Istruzioni: {istruzioni}. Modifica il contenuto seguente e restituisci SOLO il risultato. "
                                if file_name.endswith(('.xlsx', '.xls')):
                                    prompt += "Restituisci i dati in formato CSV (es: Colonna1,Colonna2\\nValore1,Valore2)."
                                else:
                                    prompt += "Restituisci solo il testo aggiornato."
                                prompt += f"\n\nContenuto:\n{contenuto}"
                                
                                testo_nuovo = ask_openai(chat_id, prompt)
                                if file_name.endswith(('.docx', '.xlsx', '.xls', '.pdf')):
                                    crea_e_invia_file_modificato(chat_id, testo_nuovo, file_name)
                                    send_message(chat_id, "✅ File modificato inviato!")
                                else:
                                    send_message(chat_id, testo_nuovo)
                            else:
                                send_message(chat_id, "📥 Analizzo il file...")
                                reply = ask_openai(chat_id, f"Analizza questo file ({file_name}) e crea un report:\n\n{contenuto}")
                                send_message(chat_id, reply)
                        elif "text" in msg:
                            reply = ask_openai(chat_id, msg["text"])
                            send_message(chat_id, reply)
        except: time.sleep(3)

if __name__ == "__main__":
    main()